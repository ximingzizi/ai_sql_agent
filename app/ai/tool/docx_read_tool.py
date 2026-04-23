import os
import json
from langchain.tools import tool
from pydantic import BaseModel, Field
from docx import Document
import pandas as pd
try:
    import pdfplumber
except:
    pdfplumber = None

class Args(BaseModel):
    path: str = Field(..., description="文件路径")

@tool("file_read_tool", args_schema=Args)
def file_read_tool(path: str) -> str:
    """
    通用文件读取工具，支持 txt、md、json、csv、docx、pdf 等文本类文件
    """
    if not os.path.exists(path):
        return f"文件不存在: {path}"

    ext = os.path.splitext(path)[-1].lower()

    try:
        # ================= TXT / MD =================
        if ext in [".txt", ".md", ".py", ".log"]:
            with open(path, "r", encoding="utf-8", errors="ignore") as f:
                return f.read()

        # ================= JSON =================
        elif ext == ".json":
            with open(path, "r", encoding="utf-8") as f:
                data = json.load(f)
            return json.dumps(data, ensure_ascii=False, indent=2)

        # ================= CSV =================
        elif ext == ".csv":

            df = pd.read_csv(path)
            return df.to_string()

        # ================= DOCX =================
        elif ext == ".docx":
            doc = Document(path)
            content = []

            # 段落
            for para in doc.paragraphs:
                content.append(para.text)

            # 表格
            for table in doc.tables:
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    content.append(" | ".join(row_data))

            return "\n".join(content)

        # ================= PDF =================
        elif ext == ".pdf":
            if not pdfplumber:
                return "未安装 pdfplumber，无法读取 PDF"

            text = []
            with pdfplumber.open(path) as pdf:
                for page in pdf.pages:
                    text.append(page.extract_text() or "")
            return "\n".join(text)

        # ==================================
        else:
            # 尝试按文本读取
            with open(path, "r", encoding="utf-8", errors="ignore") as f:
                return f.read()

    except Exception as e:
        return f"读取失败: {str(e)}"
if __name__ == "__main__":
    print(file_read_tool.invoke({"path":r"D:\PROJECT\AGENT_APP\app\static\upload\bc93b020-8ffc-46bd-9eaf-fbf4fbdf0f5a"}))