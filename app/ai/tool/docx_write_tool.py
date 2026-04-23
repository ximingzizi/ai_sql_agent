import json
import time
from pathlib import Path

from langchain.tools import tool
from pydantic import BaseModel, Field
from docx import Document

class Args(BaseModel):
    content: str = Field(..., description="写入内容（字符串或JSON字符串）")
    file_type: str = Field(..., description="文件类型: txt/md/json/csv/docx等")
    file_name: str = Field(None, description="文件名（可选）")


@tool("file_write_tool", args_schema=Args)
def file_write_tool(content: str, file_type: str, file_name: str = None) -> str:
    """
    通用文件写入工具，支持 txt、md、json、csv、docx
    """
    # 工具调用时 file_type 经常带有大小写或前导点，先统一标准化，减少模型调用失败概率。
    normalized_file_type = file_type.strip().lower().lstrip(".")

    if not file_name:
        file_name = time.strftime("%Y%m%d%H%M%S", time.localtime())

    BASE_DIR = Path(__file__).resolve().parent.parent.parent
    save_dir = BASE_DIR / "static" / "download"
    # download 目录在首次部署时可能不存在，递归创建更稳妥。
    save_dir.mkdir(parents=True, exist_ok=True)

    file_path = save_dir / f"{file_name}.{normalized_file_type}"
    print(file_path)

    try:
        if normalized_file_type in ["txt", "md"]:
            with open(file_path, "w", encoding="utf-8") as f:
                f.write(content)

        elif normalized_file_type == "json":
            data = json.loads(content)
            with open(file_path, "w", encoding="utf-8") as f:
                json.dump(data, f, ensure_ascii=False, indent=2)

        elif normalized_file_type == "csv":
            import pandas as pd
            data = json.loads(content)
            df = pd.DataFrame(data)
            df.to_csv(file_path, index=False)

        elif normalized_file_type == "docx":
            doc = Document()
            doc.add_heading("分析报告", level=1)
            for line in content.splitlines():
                if line.strip():
                    doc.add_paragraph(line)
            doc.save(file_path)

        else:
            return f"不支持的文件类型: {normalized_file_type}"

        down_path = f"http://localhost:8000/static/download/{file_name}.{normalized_file_type}"
        return f"写入成功，下载路径: {down_path}"

    except Exception as e:
        return f"写入失败: {str(e)}"
    
if __name__ == "__main__":
    from docx import Document

    # 创建Word对象
    doc = Document()

    # 写入标题
    doc.add_heading('AI项目报告', level=1)

    # 写入段落
    doc.add_paragraph("本项目是一个基于大模型的智能问答系统。")

    doc.add_paragraph("系统采用 LangChain + Agent 架构。")

    # 保存文件
    doc.save("report.docx")
